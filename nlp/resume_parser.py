import os
import pdfplumber
import docx


def extract_text_from_resume(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume not found: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _from_pdf(file_path)
    elif ext == ".docx":
        return _from_docx(file_path)
    raise ValueError(f"Unsupported format: {ext}")


def _from_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def _from_docx(path: str) -> str:
    doc   = docx.Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)
